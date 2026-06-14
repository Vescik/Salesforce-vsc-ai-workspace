"""DOCX converter (requires ``python-docx``)."""

from __future__ import annotations

from pathlib import Path
from typing import Any

from ai_workspace.knowledge.converters import make_doc


def convert(path: Path | str) -> dict[str, Any]:
    source = Path(path)
    try:
        from docx import Document  # type: ignore
    except ImportError:
        return make_doc(
            "docx",
            source,
            warnings=["python-docx not installed; .docx skipped. Install salesforce-ai-workspace-tools[knowledge]."],
            parse_status="failed",
            degraded=True,
        )

    try:
        document = Document(str(source))
    except Exception as exc:  # noqa: BLE001 - python-docx raises a variety of types
        return make_doc("docx", source, warnings=[f"python-docx failed to open file: {exc}"], parse_status="failed")

    sections: list[dict[str, Any]] = []
    tables: list[dict[str, Any]] = []
    current_heading = source.stem
    current_level = 1
    current_body: list[str] = []

    def flush() -> None:
        body = "\n".join(current_body).strip()
        if body or current_heading:
            sections.append({
                "heading": current_heading,
                "level": current_level,
                "body": body,
                "anchors": [],
            })

    for paragraph in document.paragraphs:
        style = (paragraph.style.name or "") if paragraph.style is not None else ""
        text = paragraph.text or ""
        if style.startswith("Heading"):
            flush()
            current_body = []
            try:
                current_level = int(style.split()[-1])
            except (ValueError, IndexError):
                current_level = 1
            current_heading = text.strip() or source.stem
            continue
        if text.strip():
            current_body.append(text)
    flush()

    for table in document.tables:
        rows: list[list[str]] = []
        for row in table.rows:
            rows.append([cell.text.strip() for cell in row.cells])
        if rows:
            tables.append({"caption": "", "rows": rows})

    props = document.core_properties
    metadata = {
        "author": str(props.author or ""),
        "title": str(props.title or ""),
        "created": props.created.isoformat() if props.created else "",
        "modified": props.modified.isoformat() if props.modified else "",
        "library": "python-docx",
    }

    return make_doc(
        "docx",
        source,
        sections=sections,
        tables=tables,
        metadata=metadata,
        parse_status="ok",
        degraded=False,
    )
